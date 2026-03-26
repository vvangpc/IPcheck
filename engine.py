import re
from docx import Document

class PatentEngine:
    def __init__(self, docx_path):
        self.doc = Document(docx_path)
        self.full_text = "\n".join([p.text for p in self.doc.paragraphs])
        self.sections = self._split_sections()
        self.errors = []

    def _split_sections(self):
        """将文档粗略划分为各个组成部分"""
        sections = {
            "abstract": "",
            "claims": "",
            "description": "",
            "drawings_desc": "",
            "detailed_desc": ""
        }
        
        # 简单的基于关键词的切分逻辑（实际专利文档可能有更复杂的标题）
        text = self.full_text
        
        # 寻找“权利要求书”
        claims_match = re.search(r"权利要求书", text)
        # 寻找“说明书”
        desc_match = re.search(r"说明书\n", text) or re.search(r"说明书$", text, re.M)
        
        # 这里仅作初步演示，后续可强化
        if claims_match and desc_match:
            sections["claims"] = text[claims_match.start():desc_match.start()]
            sections["description"] = text[desc_match.start():]
            
        return sections

    def check_claims_punctuation(self):
        """模块一：标点校验 - 权利要求必须以且仅以句号结尾"""
        claims_text = self.sections.get("claims", "")
        # 匹配如 "1. 一种...。" 的项
        items = re.findall(r"(\d+\..+?)(?=\n\d+\.|\s*$)", claims_text, re.S)
        
        for item in items:
            item = item.strip()
            num_match = re.match(r"(\d+)\.", item)
            if num_match:
                num = num_match.group(1)
                # 检查结尾
                if not item.endswith("。") and not item.endswith("."):
                    self.errors.append({
                        "module": "权利要求",
                        "type": "标点错误",
                        "target": f"第 {num} 项",
                        "detail": "权利要求项未以句号结尾。"
                    })
                # 检查中间是否有句号
                if item[:-1].count("。") > 0:
                     self.errors.append({
                        "module": "权利要求",
                        "type": "标点错误",
                        "target": f"第 {num} 项",
                        "detail": "权利要求项中间出现了句号（应仅在结尾出现）。"
                    })

    def check_claims_references(self):
        """模块一：引用逻辑校验 (简单演示)"""
        claims_text = self.sections.get("claims", "")
        # 提取每一项引用的数字
        # 例如："如权利要求1所述的..."
        items = re.findall(r"(\d+)\.(.+?)(?=\n\d+\.|\s*$)", claims_text, re.S)
        
        claim_map = {}
        for num, content in items:
            refs = re.findall(r"权利要求(\d+)", content)
            refs = [int(r) for r in refs]
            claim_map[int(num)] = refs
            
            # 基础校验：向后引用
            for ref in refs:
                if ref >= int(num):
                    self.errors.append({
                        "module": "权利要求",
                        "type": "引用逻辑",
                        "target": f"第 {num} 项",
                        "detail": f"存在向后引用：引用了第 {ref} 项。"
                    })

    def check_consistency(self):
        """模块二：全文一致性与对应检查"""
        # 演示：名称一致性 (提取第一行作为名称对比)
        lines = [p.text.strip() for p in self.doc.paragraphs if p.text.strip()]
        if len(lines) > 0:
            name = lines[0]
            # 检查其他地方是否包含该名称（示例逻辑）
            pass

    def check_drawing_labels(self):
        """模块二：附图标记校验 (部件名称+数字)"""
        # 匹配模式：中文名称 + 数字，如 "支架1" 或 "手柄101"
        # 排除可能是纯数字的情况
        pattern = r"([\u4e00-\u9fa5]{2,})(\d+)"
        matches = re.findall(pattern, self.full_text)
        
        name_to_id = {}
        id_to_name = {}
        
        for name, label in matches:
            # 同名不同号
            if name in name_to_id and name_to_id[name] != label:
                self.errors.append({
                    "module": "说明书",
                    "type": "附图标记",
                    "target": name,
                    "detail": f"同名不同号：在文中对应了编号 {name_to_id[name]} 和 {label}。"
                })
            # 同号不同名
            if label in id_to_name and id_to_name[label] != name:
                self.errors.append({
                    "module": "说明书",
                    "type": "附图标记",
                    "target": f"编号 {label}",
                    "detail": f"同号不同名：对应了名称 '{id_to_name[label]}' 和 '{name}'。"
                })
            
            name_to_id[name] = label
            id_to_name[label] = name

    def run_all(self):
        self.check_claims_punctuation()
        self.check_claims_references()
        self.check_drawing_labels()
        return self.errors

if __name__ == "__main__":
    # 测试代码
    pass
