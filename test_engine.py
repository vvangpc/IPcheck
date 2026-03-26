from docx import Document
from engine import PatentEngine
import os

def create_test_docx(path):
    doc = Document()
    doc.add_heading('本发明名称：一种智能支架', 0)
    
    # 模拟权利要求书
    doc.add_paragraph('权利要求书')
    doc.add_paragraph('1. 一种智能支架，其特征在于，包括底座1和支承臂2。') # 正常
    doc.add_paragraph('2. 根据权利要求1所述的智能支架，其特征在于，还包括摄像头3') # 错误：缺少句号
    doc.add_paragraph('3. 根据权利要求4所述的智能支架，其特征在于，摄像头3具有旋转功能。') # 错误：向后引用 (引用4)
    
    # 模拟说明书
    doc.add_paragraph('说明书')
    doc.add_paragraph('一种智能支架')
    doc.add_paragraph('底座1用于支撑。')
    doc.add_paragraph('支承臂100连接底座。') # 错误：同名(底座)不同号(1 vs 100)
    doc.add_paragraph('摄像头2用于拍摄。') # 错误：同号(2)不同名(支承臂 vs 摄像头)
    
    doc.save(path)

def test_engine():
    test_path = "test_patent.docx"
    create_test_docx(test_path)
    
    engine = PatentEngine(test_path)
    errors = engine.run_all()
    
    print(f"检测到 {len(errors)} 个错误:")
    for err in errors:
        print(f"[{err['module']}] {err['type']} ({err['target']}): {err['detail']}")
    
    # 清理
    if os.path.exists(test_path):
        os.remove(test_path)
    
    return len(errors) > 0

if __name__ == "__main__":
    test_engine()
